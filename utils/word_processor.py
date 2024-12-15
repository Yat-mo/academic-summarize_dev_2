import io
from docx import Document
from typing import List, Dict, Tuple
from .file_processor import BaseFileProcessor
from .exceptions import (
    FileProcessError,
    FileCorruptedError,
    TextExtractionError,
    TextProcessError
)
from config import PDFConfig

class WordProcessor(BaseFileProcessor):
    """Word文件处理器"""
    
    def _extract_text_from_file(self, file) -> str:
        """从Word文件中提取文本"""
        try:
            # 读取文件内容
            file_content = file.read()
            
            # 读取文档
            try:
                doc = Document(io.BytesIO(file_content))
            except Exception as e:
                raise FileCorruptedError(f"Word文件格式错误: {str(e)}")
            
            # 提取所有内容
            content = self._extract_document_content(doc)
            
            # 合并并清理文本
            return self._merge_content(content)
            
        except FileProcessError:
            raise
        except Exception as e:
            raise TextExtractionError(f"Word文件处理失败: {str(e)}")
    
    def _extract_document_content(self, doc: Document) -> List[Tuple[str, bool]]:
        """提取文档内容，返回(文本, 是否是标题)的列表"""
        try:
            content = []
            
            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    try:
                        # 判断是否是标题
                        is_title = paragraph.style.name.startswith('Heading')
                        # 清理文本
                        text = self._clean_text(text)
                        if text:
                            content.append((text, is_title))
                    except Exception as e:
                        raise TextProcessError(f"段落处理失败: {str(e)}")
            
            # 处理表格
            for table in doc.tables:
                try:
                    table_content = self._extract_table_content(table)
                    if table_content:
                        content.append((table_content, False))
                except Exception as e:
                    raise TextProcessError(f"表格处理失败: {str(e)}")
            
            if not content:
                raise TextExtractionError("文档内容为空")
            
            return content
            
        except FileProcessError:
            raise
        except Exception as e:
            raise TextExtractionError(f"文档内容提取失败: {str(e)}")
    
    def _extract_table_content(self, table) -> str:
        """提取表格内容"""
        try:
            rows_content = []
            
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        try:
                            # 清理文本
                            text = self._clean_text(text)
                            if text:
                                row_texts.append(text)
                        except Exception as e:
                            raise TextProcessError(f"单元格处理失败: {str(e)}")
                
                if row_texts:
                    rows_content.append(' | '.join(row_texts))
            
            return '\n'.join(rows_content) if rows_content else ""
            
        except FileProcessError:
            raise
        except Exception as e:
            raise TextExtractionError(f"表格内容提取失败: {str(e)}")
    
    def _merge_content(self, content: List[Tuple[str, bool]]) -> str:
        """合并文档内容，保持标题层级关系"""
        try:
            if not content:
                raise TextProcessError("没有要合并的内容")
                
            merged = []
            current_title = None
            
            for text, is_title in content:
                if is_title:
                    current_title = text
                    merged.append(text)
                else:
                    # 如果有标题且不是第一个段落，添加标题上下文
                    if current_title and merged:
                        text = f"{current_title}\n{text}"
                    merged.append(text)
            
            if not merged:
                raise TextProcessError("合并后的内容为空")
            
            return '\n\n'.join(merged)
            
        except FileProcessError:
            raise
        except Exception as e:
            raise TextExtractionError(f"内容合并失败: {str(e)}")
